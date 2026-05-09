#!/usr/bin/env python3
"""
Genera un documento Word con todos los LABs de CryptoVulnX y su resolucion paso a paso.

Lee labs/LAB01-*.md hasta LAB14-*.md y arma un docx con:
- Portada
- Indice
- Introduccion a la metodologia
- Por cada LAB: titulo, contenido convertido de markdown
"""

import re
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
LABS_DIR = ROOT / "labs"
OUT = ROOT / "CryptoVulnX-Labs-Completo.docx"

# Orden de los labs en el documento
LAB_ORDER = [
    ("LAB11-RECON.md",                        "Fase 1 - Metodologia"),
    ("LAB12-API-INVENTORY-FUZZING.md",        "Fase 2 - Metodologia"),
    ("LAB13-ENDPOINT-FUZZING.md",             "Fase 3 - Metodologia"),
    ("LAB14-PARAMETER-FUZZING.md",            "Fase 4 - Metodologia"),
    ("LAB01-BOLA.md",                         "Fase 5 - OWASP API Top 10"),
    ("LAB02-BROKEN-AUTH.md",                  "Fase 5 - OWASP API Top 10"),
    ("LAB03-BOPLA.md",                        "Fase 5 - OWASP API Top 10"),
    ("LAB04-RESOURCE.md",                     "Fase 5 - OWASP API Top 10"),
    ("LAB05-BFLA.md",                         "Fase 5 - OWASP API Top 10"),
    ("LAB06-BUSINESS-FLOW.md",                "Fase 5 - OWASP API Top 10"),
    ("LAB07-SSRF.md",                         "Fase 5 - OWASP API Top 10"),
    ("LAB08-MISCONFIG.md",                    "Fase 5 - OWASP API Top 10"),
    ("LAB09-INVENTORY.md",                    "Fase 5 - OWASP API Top 10"),
    ("LAB10-UNSAFE-CONSUMPTION.md",           "Fase 5 - OWASP API Top 10"),
]

CODE_BG = "F4F4F4"
TABLE_HEADER_BG = "2C3E50"


def add_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_paragraph_shading(paragraph, color_hex):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    p_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def setup_styles(doc):
    styles = doc.styles
    # Code block style
    if "CodeBlock" not in styles:
        cs = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        cs.font.name = "Consolas"
        cs.font.size = Pt(9)
        cs.paragraph_format.left_indent = Cm(0.5)
        cs.paragraph_format.space_before = Pt(4)
        cs.paragraph_format.space_after = Pt(4)
    # Inline code
    if "InlineCode" not in styles:
        ic = styles.add_style("InlineCode", WD_STYLE_TYPE.CHARACTER)
        ic.font.name = "Consolas"
        ic.font.size = Pt(10)


# ----------- Markdown parser muy simple -----------

CODE_FENCE_RE = re.compile(r"^```(\w*)\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
TABLE_SEP_RE = re.compile(r"^\|?\s*(:?-+:?\s*\|)+\s*:?-+:?\s*\|?\s*$")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
HR_RE = re.compile(r"^\s*---+\s*$")
LIST_RE = re.compile(r"^(\s*)([-*]|\d+\.)\s+(.+)$")
DETAILS_OPEN_RE = re.compile(r"^<details>")
DETAILS_CLOSE_RE = re.compile(r"^</details>")
SUMMARY_RE = re.compile(r"^<summary>(.+)</summary>")


def parse_inline(run_target_paragraph, text):
    """Maneja **bold** y `code` inline. run_target_paragraph es un objeto paragraph."""
    # Patron: alterna texto / **bold** / `code`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = run_target_paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = run_target_paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run_target_paragraph.add_run(part)


def render_table(doc, header_cells, rows):
    if not header_cells:
        return
    n_cols = len(header_cells)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    # Header
    for i, h in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        add_shading(cell, TABLE_HEADER_BG)
    # Rows
    for r_idx, row in enumerate(rows, start=1):
        cells = row + [""] * (n_cols - len(row))
        for c_idx in range(n_cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            parse_inline(p, cells[c_idx].strip())


def render_code_block(doc, lines, lang=""):
    p = doc.add_paragraph(style="CodeBlock")
    add_paragraph_shading(p, CODE_BG)
    if lang:
        run = p.add_run(f"[{lang}]\n")
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def render_markdown(doc, md_text, base_heading_level=2):
    """Parser de markdown simple. base_heading_level=2 -> # del MD se vuelve Heading 2 en Word."""
    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # Code fence
        m = CODE_FENCE_RE.match(line)
        if m:
            lang = m.group(1) or ""
            i += 1
            code_lines = []
            while i < n and not CODE_FENCE_RE.match(lines[i]):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            render_code_block(doc, code_lines, lang)
            continue

        # Details (collapsible) - lo expandimos como sub-seccion plana
        if DETAILS_OPEN_RE.match(line):
            i += 1
            # Buscar summary
            if i < n:
                m_sum = SUMMARY_RE.match(lines[i])
                if m_sum:
                    p = doc.add_paragraph()
                    run = p.add_run(m_sum.group(1).strip())
                    run.bold = True
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x44, 0x66, 0x88)
                    i += 1
            # Procesar contenido recursivamente hasta </details>
            inner_lines = []
            while i < n and not DETAILS_CLOSE_RE.match(lines[i]):
                inner_lines.append(lines[i])
                i += 1
            i += 1  # consume </details>
            render_markdown(doc, "\n".join(inner_lines), base_heading_level)
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading_level = min(9, base_heading_level + level - 1)
            p = doc.add_heading(text, level=heading_level)
            i += 1
            continue

        # Table
        if line.startswith("|") and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < n and lines[i].startswith("|") and lines[i].strip().endswith("|"):
                row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row_cells)
                i += 1
            render_table(doc, header, rows)
            continue

        # Horizontal rule
        if HR_RE.match(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("―" * 40)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            i += 1
            continue

        # List
        m = LIST_RE.match(line)
        if m:
            indent = len(m.group(1))
            text = m.group(3)
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                p.paragraph_format.left_indent = Cm(1.5)
            parse_inline(p, text)
            i += 1
            continue

        # Paragraph (collect contiguous non-empty lines)
        if line.strip() == "":
            i += 1
            continue

        para_lines = []
        while i < n and lines[i].strip() != "" and not (
            CODE_FENCE_RE.match(lines[i]) or HEADING_RE.match(lines[i])
            or HR_RE.match(lines[i]) or LIST_RE.match(lines[i])
            or lines[i].startswith("|")
            or DETAILS_OPEN_RE.match(lines[i]) or DETAILS_CLOSE_RE.match(lines[i])
        ):
            para_lines.append(lines[i])
            i += 1
        text = " ".join(para_lines).strip()
        if text:
            p = doc.add_paragraph()
            parse_inline(p, text)


# ----------- Documento -----------

def build_document():
    doc = Document()
    setup_styles(doc)

    # ---- Portada ----
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CryptoVulnX")
    run.font.size = Pt(48)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Laboratorio completo de pentest de APIs")
    run.font.size = Pt(20)
    run.italic = True

    doc.add_paragraph()
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Resolucion paso a paso de los 14 LABs")
    run.font.size = Pt(16)

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub3.add_run("Metodologia 5 fases + OWASP API Security Top 10 (2023)")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(8):
        doc.add_paragraph()

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run("Documento educativo - uso exclusivo en entornos aislados")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    add_page_break(doc)

    # ---- Indice ----
    doc.add_heading("Indice", level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "Este documento contiene los 14 LABs de CryptoVulnX en su orden metodologico. "
        "Los LABs 11-14 introducen las fases de descubrimiento (RECON, INVENTORY, "
        "ENDPOINT FUZZING, PARAMETER FUZZING). Los LABs 01-10 cubren la fase de "
        "explotacion contra el OWASP API Security Top 10 (2023)."
    )

    doc.add_paragraph()
    current_section = None
    for i, (md_file, section) in enumerate(LAB_ORDER, start=1):
        if section != current_section:
            p = doc.add_paragraph()
            run = p.add_run(section)
            run.bold = True
            run.font.size = Pt(12)
            current_section = section
        title = md_file.replace(".md", "")
        p = doc.add_paragraph(style="List Number")
        p.add_run(title)

    add_page_break(doc)

    # ---- Introduccion a la metodologia ----
    doc.add_heading("Introduccion a la metodologia", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "El pentest de un API moderno NO es ejecutar SQLi contra /api/v1/login. "
        "Es un proceso ordenado de descubrimiento y explotacion donde cada fase "
        "alimenta a la siguiente. El error frecuente del alumno principiante es "
        "saltar directo a la fase 5. El pentest profesional gasta el "
    )
    run = p.add_run("70% del tiempo en fases 1-4")
    run.bold = True
    p.add_run(
        " y solo el 30% en explotacion. Razon: si la enumeracion fue completa, "
        "la explotacion es trivial."
    )

    doc.add_heading("Las 5 fases", level=2)
    fases = [
        ("Fase 1 - RECON",
         "Descubrir archivos publicos: robots.txt, .git/, .env*, backups, notes, dumps SQL."),
        ("Fase 2 - API INVENTORY",
         "Mapear todas las versiones del API: v1, v2, v3, internal, test, dev, staging. Buscar swagger.json y openapi.json."),
        ("Fase 3 - ENDPOINT FUZZING",
         "Brute paths bajo cada version. Verb fuzzing. Hallazgos tipicos: admin/exec, admin/backup, internal/health."),
        ("Fase 4 - PARAMETER FUZZING",
         "Descubrir query params, body params, headers y cookies no documentados que cambian comportamiento."),
        ("Fase 5 - EXPLOITATION",
         "Combinar todo lo descubierto en fases 1-4 para ejecutar chains de OWASP API Top 10.")
    ]
    for nombre, desc in fases:
        p = doc.add_paragraph()
        run = p.add_run(nombre)
        run.bold = True
        p.add_run(": " + desc)

    doc.add_heading("Como leer cada LAB", level=2)
    for txt in [
        "Cada LAB tiene la misma estructura: Objetivo, Contexto, Endpoints involucrados, Dificultad, Pistas, Solucion paso a paso, Remediacion.",
        "Los bloques de codigo (Consolas, fondo gris) son comandos listos para copiar/pegar.",
        "Los recuadros 'Pista N' son guias progresivas: empezar por la primera y avanzar solo si te trabas.",
        "La seccion 'Solucion' es la respuesta. No leerla antes de intentar la fase.",
        "La seccion 'Remediacion' muestra el codigo correcto - util como referencia para el equipo de desarrollo."
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(txt)

    add_page_break(doc)

    # ---- Cuerpo: cada LAB ----
    current_section = None
    for idx, (md_file, section) in enumerate(LAB_ORDER, start=1):
        path = LABS_DIR / md_file
        if not path.exists():
            print(f"WARN: {path} no existe, saltando")
            continue

        # Section divider
        if section != current_section:
            doc.add_heading(section, level=1)
            current_section = section

        text = path.read_text(encoding="utf-8")
        # El primer heading del MD se vuelve heading 1, los demas se desplazan
        render_markdown(doc, text, base_heading_level=2)

        add_page_break(doc)

    # ---- Apendice: Resumen final ----
    doc.add_heading("Apendice - Resumen del flujo completo", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "El recorrido recomendado para resolver todos los LABs en orden es:\n"
        "(1) Levantar el lab con docker compose up -d. "
        "(2) Acceder al playbook gamificado en /playbook.php para guiarte. "
        "(3) Resolver cada fase y conseguir flags. "
        "(4) Volver a este documento como referencia de soluciones cuando te trabes."
    )

    doc.add_heading("Encadenamiento entre fases", level=2)
    chain_table = [
        ["Hallazgo de Fase 1", "Habilita en Fase..."],
        ["JWT_SECRET=crypto123 (en notes.txt, .env.bak, debug.php)", "LAB02 - forjar JWT con role=admin"],
        ["admin/admin123 en backup.sql", "LAB02, LAB05 - login directo como admin"],
        ["Lista de endpoints en notes.txt y openapi.json", "LAB12, LAB13 - fuzzing dirigido"],
        ["Magic headers en notes.txt", "LAB14 - header fuzzing"],
        ["composer.lock con CVEs", "LAB08 - explotacion de libreria vulnerable"],
        ["RCE en /api/v3/admin/exec.php", "Lectura de .env real, dump de DB"],
        ["X-Admin-Token en magic params", "LAB05 BFLA sin necesidad de JWT admin"],
        ["rate_override, bypass_kyc, fee_override", "LAB06 Business Flow"],
        ["X-Forwarded-For, X-Original-URL", "LAB07 SSRF"],
    ]
    render_table(doc, chain_table[0], chain_table[1:])

    # Save
    doc.save(OUT)
    print(f"OK: documento generado en {OUT}")
    print(f"   Tamano: {OUT.stat().st_size // 1024} KB")


if __name__ == "__main__":
    try:
        build_document()
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
